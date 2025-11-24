import pandas as pd
import numpy as np
from scipy import stats
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings

# Suppress warnings
warnings.filterwarnings("ignore")

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for k, v in kwargs.items():
        tag = 'w:{}'.format(k)
        element = tcPr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcPr.append(element)

        for attr_key, attr_value in v.items():
            element.set(qn('w:{}'.format(attr_key)), str(attr_value))

def create_table_one(csv_path, output_path):
    # Load data
    try:
        df = pd.read_csv(csv_path)
    except Exception as e:
        print(f"Error reading CSV: {e}")
        return

    # Define variables
    group_col = 'edm_18'
    if group_col not in df.columns:
        if 'edr_18m' in df.columns:
            print(f"Warning: '{group_col}' not found. Using 'edr_18m' instead.")
            group_col = 'edr_18m'
        else:
            print(f"Error: Grouping column '{group_col}' not found.")
            return

    # Drop rows where group_col is NaN
    df = df.dropna(subset=[group_col])
    
    # Define variable lists
    continuous_vars = [
        'Age', 'BMI', 'LN_Total', 'LN_Positive', 'LNR', 
        'Tumor_Size_cm', 'CEA_PreOp', 'PreOp_Albumin'
    ]
    
    categorical_vars = [
        'Sex', 'ECOG', 'Tumor_Location_Group', 'pT_Stage', 'pN_Stage', 
        'AJCC_Substage', 'Histology', 'Differentiation', 'LVI', 'PNI', 
        'MSI_High', 'Death'
    ]

    # Map for display names if needed (optional, using raw names for now or simple cleanup)
    
    # Initialize DOCX
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # Title
    heading = doc.add_heading('Table 1. Baseline Characteristics', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create table
    # Columns: Variable, Total (N=...), Group 0 (N=...), Group 1 (N=...), p-value
    groups = sorted(df[group_col].unique())
    if len(groups) != 2:
        print(f"Warning: Grouping variable has {len(groups)} levels. Expecting 2 for binary comparison.")
        # Continue but p-value logic might need adjustment if > 2 groups (ANOVA/Kruskal)
        # The prompt implies 2 groups (t-test/Mann-Whitney).
        # We will proceed assuming 2 groups for p-value, or pairwise if more?
        # Let's assume 2 groups as per "baseline characteristics" typical usage.
    
    group0_label = f"No ({group_col}=0)"
    group1_label = f"Yes ({group_col}=1)"
    # Adjust labels based on actual values if they are not 0/1
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = f'Total (N={len(df)})'
    
    # Split data
    group_data = {}
    for g in groups:
        group_data[g] = df[df[group_col] == g]
    
    # Set headers for groups
    for i, g in enumerate(groups):
        hdr_cells[i+2].text = f'Group {g} (N={len(group_data[g])})'
    
    hdr_cells[4].text = 'p-value'

    # Helper for p-value formatting
    def format_p(p):
        if pd.isna(p):
            return "N/A"
        
        val_str = ""
        if p < 0.001:
            val_str = "<0.001"
        else:
            val_str = f"{p:.3f}"
            
        if p < 0.05:
            val_str += "*"
        return val_str

    # Process Continuous Variables
    for var in continuous_vars:
        if var not in df.columns:
            continue
            
        row_cells = table.add_row().cells
        row_cells[0].text = var
        
        # Check normality on the whole dataset or per group? 
        # Usually check per group for t-test assumptions, or check residuals. 
        # Prompt says: "如果是continuous numeric，先用Shapiro-Wilk進行常態檢定"
        # We will check normality for EACH group. If BOTH are normal -> t-test. Else -> Mann-Whitney.
        
        is_normal = True
        group_stats = {}
        
        for g in groups:
            data_g = group_data[g][var].dropna()
            if len(data_g) < 3: # Shapiro requires at least 3
                is_normal = False # Fallback to non-parametric if too small
            else:
                stat, p_shapiro = stats.shapiro(data_g)
                if p_shapiro < 0.05:
                    is_normal = False
            
            group_stats[g] = data_g

        # Calculate Total stats
        data_total = df[var].dropna()
        
        # Decide presentation based on normality
        # If normal: Mean +/- SD
        # If non-normal: Median (IQR)
        
        # Note: Usually we stick to one format for the row. 
        # If ANY group is non-normal, we usually use Median (IQR) and MWU.
        
        if is_normal:
            # Parametric
            # Total
            mean_tot = data_total.mean()
            sd_tot = data_total.std()
            row_cells[1].text = f"{mean_tot:.1f} ± {sd_tot:.1f}"
            
            # Groups
            group_vals = []
            for i, g in enumerate(groups):
                d = group_stats[g]
                m = d.mean()
                s = d.std()
                row_cells[i+2].text = f"{m:.1f} ± {s:.1f}"
                group_vals.append(d)
            
            # Test
            if len(groups) == 2:
                t_stat, p_val = stats.ttest_ind(group_vals[0], group_vals[1], equal_var=False) # Welch's t-test is safer
            else:
                p_val = np.nan # Not handling >2 groups yet
                
        else:
            # Non-parametric
            # Total
            med_tot = data_total.median()
            q1_tot = data_total.quantile(0.25)
            q3_tot = data_total.quantile(0.75)
            row_cells[1].text = f"{med_tot:.1f} ({q1_tot:.1f}-{q3_tot:.1f})"
            
            # Groups
            group_vals = []
            for i, g in enumerate(groups):
                d = group_stats[g]
                med = d.median()
                q1 = d.quantile(0.25)
                q3 = d.quantile(0.75)
                row_cells[i+2].text = f"{med:.1f} ({q1:.1f}-{q3:.1f})"
                group_vals.append(d)
            
            # Test
            if len(groups) == 2:
                u_stat, p_val = stats.mannwhitneyu(group_vals[0], group_vals[1])
            else:
                p_val = np.nan

        row_cells[4].text = format_p(p_val)

    # Process Categorical Variables
    for var in categorical_vars:
        if var not in df.columns:
            continue
            
        # Add variable name row
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[0].merge(row_cells[4]) # Merge header row for categorical var
        
        # Filter out missing for stats and valid counts
        valid_data = df.dropna(subset=[var])
        
        # Get unique categories (excluding NaN)
        categories = sorted(valid_data[var].unique())
        
        # Create contingency table for testing (using only valid data)
        contingency_table = pd.crosstab(valid_data[var], valid_data[group_col])
        
        # Check expected frequencies
        if contingency_table.size > 0:
            chi2, p_chi2, dof, expected = stats.chi2_contingency(contingency_table)
            
            use_fisher = False
            if (expected < 5).any():
                # If 2x2, use Fisher
                if contingency_table.shape == (2, 2):
                    use_fisher = True
                    oddsr, p_val = stats.fisher_exact(contingency_table, alternative='two-sided')
                else:
                    p_val = p_chi2
            else:
                p_val = p_chi2
        else:
            p_val = np.nan

        # Calculate valid counts for denominators
        n_valid_total = len(valid_data)
        n_valid_groups = {}
        for g in groups:
            n_valid_groups[g] = len(valid_data[valid_data[group_col] == g])

        # Add rows for each category
        first_cat = True
        for cat in categories:
            row_cells = table.add_row().cells
            row_cells[0].text = f"  {cat}" # Indent
            
            # Total
            n_tot = len(valid_data[valid_data[var] == cat])
            pct_tot = (n_tot / n_valid_total) * 100 if n_valid_total > 0 else 0
            row_cells[1].text = f"{n_tot} ({pct_tot:.1f}%)"
            
            # Groups
            for i, g in enumerate(groups):
                n_g = len(valid_data[(valid_data[var] == cat) & (valid_data[group_col] == g)])
                n_group_valid = n_valid_groups.get(g, 0)
                pct_g = (n_g / n_group_valid) * 100 if n_group_valid > 0 else 0
                row_cells[i+2].text = f"{n_g} ({pct_g:.1f}%)"
            
            if first_cat:
                row_cells[4].text = format_p(p_val)
                first_cat = False
            else:
                row_cells[4].text = ""
        
        # Handle Missing Values
        n_missing_total = len(df) - n_valid_total
        if n_missing_total > 0:
            row_cells = table.add_row().cells
            row_cells[0].text = "  Missing"
            
            # Total Missing % (usually based on total N including missing, or just count)
            # User asked: "缺值拉出來成一行並計算缺值比例"
            # Usually missing % is N_missing / N_total_dataset
            pct_missing_tot = (n_missing_total / len(df)) * 100
            row_cells[1].text = f"{n_missing_total} ({pct_missing_tot:.1f}%)"
            
            for i, g in enumerate(groups):
                # Count missing in this group
                # Original group data includes missing in var
                n_g_missing = len(group_data[g]) - n_valid_groups.get(g, 0)
                pct_g_missing = (n_g_missing / len(group_data[g])) * 100
                row_cells[i+2].text = f"{n_g_missing} ({pct_g_missing:.1f}%)"
            
            row_cells[4].text = "" # No p-value for missing row

    # Remove bottom border of the last row (or whole table bottom)
    # Since we used 'Table Grid', it has borders everywhere.
    # To remove the bottom border of the table, we can modify the last row's cells.
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 0, "val": "none"})

    # Add Footnotes
    doc.add_paragraph() # Spacer
    footnote = doc.add_paragraph()
    footnote.style = 'Normal'
    run1 = footnote.add_run("Abbreviations: SD, Standard Deviation; IQR, Interquartile Range.\n")
    run1.font.size = Pt(9)
    run2 = footnote.add_run("Statistical tests: Continuous variables were tested using Student's t-test (normal distribution) or Mann-Whitney U test (non-normal distribution). Categorical variables were tested using Chi-square test or Fisher's exact test. * p < 0.05.")
    run2.font.size = Pt(9)

    # Save
    doc.save(output_path)
    print(f"Table saved to {output_path}")

if __name__ == "__main__":
    create_table_one('notebooks/data_typed.csv', 'baseline_characteristics.docx')
